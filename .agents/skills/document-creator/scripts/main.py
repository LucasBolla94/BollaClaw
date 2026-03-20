#!/usr/bin/env python3
"""
BollaClaw Document Creator — PDF, DOCX, XLSX generation
Receives JSON from stdin with __tool__ field to dispatch to the right generator.
"""

import sys
import os
import json
from datetime import datetime
from pathlib import Path

OUTPUT_DIR = Path("./output")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def create_pdf(args: dict) -> dict:
    """Create a professional PDF using ReportLab."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm, mm
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        PageBreak, HRFlowable
    )
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY

    filename = args.get("filename", "documento.pdf")
    if not filename.endswith(".pdf"):
        filename += ".pdf"

    title = args.get("title", "Documento")
    content = args.get("content", "")
    author = args.get("author", "BollaClaw")
    font_size = int(args.get("font_size", 11))
    lang = args.get("language", "pt")

    filepath = OUTPUT_DIR / filename

    # Create document
    doc = SimpleDocTemplate(
        str(filepath),
        pagesize=A4,
        topMargin=2.5 * cm,
        bottomMargin=2 * cm,
        leftMargin=2.5 * cm,
        rightMargin=2.5 * cm,
        title=title,
        author=author,
    )

    # Styles
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        "DocTitle",
        parent=styles["Title"],
        fontSize=font_size + 8,
        spaceAfter=20,
        textColor=HexColor("#1a1a2e"),
    ))
    styles.add(ParagraphStyle(
        "SectionHeader",
        parent=styles["Heading2"],
        fontSize=font_size + 3,
        spaceBefore=16,
        spaceAfter=8,
        textColor=HexColor("#16213e"),
    ))
    styles.add(ParagraphStyle(
        "SubHeader",
        parent=styles["Heading3"],
        fontSize=font_size + 1,
        spaceBefore=12,
        spaceAfter=6,
        textColor=HexColor("#0f3460"),
    ))
    styles.add(ParagraphStyle(
        "BodyText2",
        parent=styles["BodyText"],
        fontSize=font_size,
        leading=font_size + 5,
        alignment=TA_JUSTIFY,
        spaceAfter=8,
    ))
    styles.add(ParagraphStyle(
        "BulletItem",
        parent=styles["BodyText"],
        fontSize=font_size,
        leftIndent=20,
        bulletIndent=10,
        spaceAfter=4,
    ))

    elements = []

    # Title
    elements.append(Paragraph(title, styles["DocTitle"]))

    # Date and author line
    now = datetime.now()
    if lang == "pt":
        months_pt = ["Janeiro", "Fevereiro", "Março", "Abril", "Maio", "Junho",
                      "Julho", "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro"]
        date_str = f"{now.day} de {months_pt[now.month-1]} de {now.year}"
    else:
        date_str = now.strftime("%B %d, %Y")

    meta_text = f"<i>{author} — {date_str}</i>"
    elements.append(Paragraph(meta_text, styles["BodyText"]))
    elements.append(HRFlowable(width="100%", thickness=1, color=HexColor("#cccccc")))
    elements.append(Spacer(1, 12))

    # Parse content
    lines = content.split("\\n") if "\\n" in content else content.split("\n")

    table_buffer = []

    def flush_table():
        nonlocal table_buffer
        if not table_buffer:
            return
        # Parse table
        table_data = []
        for row_str in table_buffer:
            cols = [c.strip() for c in row_str.split("|") if c.strip()]
            table_data.append(cols)

        if table_data:
            # Ensure all rows same length
            max_cols = max(len(r) for r in table_data)
            table_data = [r + [""] * (max_cols - len(r)) for r in table_data]

            t = Table(table_data, repeatRows=1)
            style = TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), HexColor("#1a1a2e")),
                ("TEXTCOLOR", (0, 0), (-1, 0), HexColor("#ffffff")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, 0), font_size),
                ("FONTSIZE", (0, 1), (-1, -1), font_size - 1),
                ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#cccccc")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [HexColor("#f8f9fa"), HexColor("#ffffff")]),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
            ])
            t.setStyle(style)
            elements.append(Spacer(1, 8))
            elements.append(t)
            elements.append(Spacer(1, 8))

        table_buffer = []

    for line in lines:
        stripped = line.strip()

        if not stripped:
            flush_table()
            elements.append(Spacer(1, 6))
            continue

        # Table row (contains | separator)
        if "|" in stripped and not stripped.startswith("#") and not stripped.startswith("-"):
            table_buffer.append(stripped)
            continue
        else:
            flush_table()

        # Section header ##
        if stripped.startswith("## "):
            elements.append(Paragraph(stripped[3:], styles["SectionHeader"]))
        elif stripped.startswith("### "):
            elements.append(Paragraph(stripped[4:], styles["SubHeader"]))
        elif stripped.startswith("# "):
            elements.append(Paragraph(stripped[2:], styles["SectionHeader"]))
        # Bullet point
        elif stripped.startswith("- ") or stripped.startswith("• "):
            bullet_text = stripped[2:]
            elements.append(Paragraph(f"• {bullet_text}", styles["BulletItem"]))
        # Page break
        elif stripped == "---":
            elements.append(PageBreak())
        else:
            # Regular paragraph — handle basic bold/italic
            text = stripped
            text = text.replace("**", "<b>", 1)
            text = text.replace("**", "</b>", 1)
            elements.append(Paragraph(text, styles["BodyText2"]))

    flush_table()

    # Build PDF
    doc.build(elements)

    file_size = os.path.getsize(filepath)
    return {
        "success": True,
        "filepath": str(filepath),
        "filename": filename,
        "size_bytes": file_size,
        "message": f"PDF criado: {filepath} ({file_size} bytes)\n[FILE:{filepath}]"
    }


def create_docx(args: dict) -> dict:
    """Create a professional Word document using python-docx."""
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    filename = args.get("filename", "documento.docx")
    if not filename.endswith(".docx"):
        filename += ".docx"

    title = args.get("title", "Documento")
    content = args.get("content", "")
    author = args.get("author", "BollaClaw")
    font_name = args.get("font_name", "Calibri")
    font_size = int(args.get("font_size", 11))

    filepath = OUTPUT_DIR / filename

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(font_size)
    style.paragraph_format.space_after = Pt(6)

    # Core properties
    doc.core_properties.author = author
    doc.core_properties.title = title

    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Date
    now = datetime.now()
    date_para = doc.add_paragraph()
    run = date_para.add_run(f"{author} — {now.strftime('%d/%m/%Y')}")
    run.italic = True
    run.font.size = Pt(font_size - 1)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # Separator
    doc.add_paragraph("_" * 60)

    # Parse content
    lines = content.split("\\n") if "\\n" in content else content.split("\n")

    table_buffer = []

    def flush_table():
        nonlocal table_buffer
        if not table_buffer:
            return

        table_data = []
        for row_str in table_buffer:
            cols = [c.strip() for c in row_str.split("|") if c.strip()]
            table_data.append(cols)

        if table_data:
            max_cols = max(len(r) for r in table_data)
            table_data = [r + [""] * (max_cols - len(r)) for r in table_data]

            table = doc.add_table(rows=len(table_data), cols=max_cols)
            table.style = "Light Grid Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            for i, row_data in enumerate(table_data):
                for j, cell_val in enumerate(row_data):
                    cell = table.cell(i, j)
                    cell.text = str(cell_val)
                    # Bold headers
                    if i == 0:
                        for p in cell.paragraphs:
                            for r in p.runs:
                                r.bold = True

            doc.add_paragraph()  # Space after table

        table_buffer = []

    for line in lines:
        stripped = line.strip()

        if not stripped:
            flush_table()
            doc.add_paragraph()
            continue

        # Table row
        if "|" in stripped and not stripped.startswith("#") and not stripped.startswith("-"):
            table_buffer.append(stripped)
            continue
        else:
            flush_table()

        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("• "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped == "---":
            doc.add_page_break()
        else:
            para = doc.add_paragraph()
            # Handle bold markers
            parts = stripped.split("**")
            for idx, part in enumerate(parts):
                if part:
                    run = para.add_run(part)
                    run.font.name = font_name
                    run.font.size = Pt(font_size)
                    if idx % 2 == 1:  # Odd index = inside bold markers
                        run.bold = True

    flush_table()

    doc.save(str(filepath))

    file_size = os.path.getsize(filepath)
    return {
        "success": True,
        "filepath": str(filepath),
        "filename": filename,
        "size_bytes": file_size,
        "message": f"DOCX criado: {filepath} ({file_size} bytes)\n[FILE:{filepath}]"
    }


def create_xlsx(args: dict) -> dict:
    """Create a formatted Excel spreadsheet using openpyxl."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    filename = args.get("filename", "dados.xlsx")
    if not filename.endswith(".xlsx"):
        filename += ".xlsx"

    sheet_name = args.get("sheet_name", "Dados")
    headers = args.get("headers", [])
    rows = args.get("rows", [])
    title = args.get("title", "")
    auto_filter = args.get("auto_filter", True)

    filepath = OUTPUT_DIR / filename

    wb = Workbook()
    ws = wb.active
    ws.title = sheet_name

    # Styles
    header_font = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="1a1a2e", end_color="1a1a2e", fill_type="solid")
    header_align = Alignment(horizontal="center", vertical="center", wrap_text=True)

    data_font = Font(name="Calibri", size=10)
    data_align = Alignment(vertical="center", wrap_text=True)

    thin_border = Border(
        left=Side(style="thin", color="CCCCCC"),
        right=Side(style="thin", color="CCCCCC"),
        top=Side(style="thin", color="CCCCCC"),
        bottom=Side(style="thin", color="CCCCCC"),
    )

    alt_fill = PatternFill(start_color="F8F9FA", end_color="F8F9FA", fill_type="solid")

    current_row = 1

    # Optional title row
    if title:
        title_font = Font(name="Calibri", size=14, bold=True, color="1a1a2e")
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=max(len(headers), 1))
        title_cell = ws.cell(row=1, column=1, value=title)
        title_cell.font = title_font
        title_cell.alignment = Alignment(horizontal="left")
        current_row = 3  # Skip a row after title

    # Headers
    if headers:
        for col_idx, header in enumerate(headers, 1):
            cell = ws.cell(row=current_row, column=col_idx, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_align
            cell.border = thin_border

        # Auto-filter
        if auto_filter:
            ws.auto_filter.ref = f"A{current_row}:{get_column_letter(len(headers))}{current_row + len(rows)}"

        current_row += 1

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, value in enumerate(row_data, 1):
            # Try to convert numeric strings
            if isinstance(value, str):
                try:
                    value = float(value) if "." in value else int(value)
                except (ValueError, TypeError):
                    pass

            cell = ws.cell(row=current_row, column=col_idx, value=value)
            cell.font = data_font
            cell.alignment = data_align
            cell.border = thin_border

            # Alternate row coloring
            if row_idx % 2 == 0:
                cell.fill = alt_fill

        current_row += 1

    # Auto-width columns
    for col_idx in range(1, (len(headers) or 1) + 1):
        max_length = 0
        col_letter = get_column_letter(col_idx)

        for row in ws.iter_rows(min_col=col_idx, max_col=col_idx, values_only=False):
            for cell in row:
                if cell.value:
                    max_length = max(max_length, len(str(cell.value)))

        ws.column_dimensions[col_letter].width = min(max_length + 4, 50)

    # Freeze header row
    header_row = 3 if title else 1
    ws.freeze_panes = f"A{header_row + 1}"

    wb.save(str(filepath))

    file_size = os.path.getsize(filepath)
    total_rows = len(rows)
    total_cols = len(headers)

    return {
        "success": True,
        "filepath": str(filepath),
        "filename": filename,
        "size_bytes": file_size,
        "rows": total_rows,
        "columns": total_cols,
        "message": f"XLSX criado: {filepath} ({total_rows} linhas, {total_cols} colunas)\n[FILE:{filepath}]"
    }


def main():
    """Main dispatcher — reads JSON from stdin, routes to the correct tool."""
    try:
        raw = sys.stdin.read().strip()
        if not raw:
            print(json.dumps({"error": "No input received"}))
            return

        args = json.loads(raw)
        tool = args.pop("__tool__", "")

        if tool == "create_pdf":
            result = create_pdf(args)
        elif tool == "create_docx":
            result = create_docx(args)
        elif tool == "create_xlsx":
            result = create_xlsx(args)
        else:
            result = {"error": f"Unknown tool: {tool}. Available: create_pdf, create_docx, create_xlsx"}

        print(json.dumps(result, ensure_ascii=False, indent=2))

    except ImportError as e:
        missing = str(e).split("'")[-2] if "'" in str(e) else str(e)
        print(json.dumps({
            "error": f"Missing dependency: {missing}. Run: pip3 install reportlab python-docx openpyxl --break-system-packages",
        }))
    except Exception as e:
        print(json.dumps({"error": f"Document creation failed: {str(e)}"}))


if __name__ == "__main__":
    main()
